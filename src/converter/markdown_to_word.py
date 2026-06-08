"""Markdown → Word 转换器"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def convert_markdown_to_word(markdown: str, output_path: str, title: str = "") -> str:
    """将 Markdown 内容转换为 .docx 文件

    Args:
        markdown: Markdown 文本
        output_path: 输出 .docx 文件路径
        title: 文档标题（可选，作为第一行）

    Returns:
        输出文件路径

    Raises:
        ValueError: markdown 为空
    """
    if not markdown or not markdown.strip():
        raise ValueError("content 不能为空")

    doc = Document()

    # 标题（如果有）
    if title:
        p = doc.add_paragraph(title)
        p.runs[0].bold = True

    # 逐行解析
    for line in markdown.split("\n"):
        line = line.rstrip()

        # 空行跳过
        if not line:
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            continue

        # 无序列表
        list_match = re.match(r"^[-*+]\s+(.+)$", line)
        if list_match:
            text = list_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_formatted_text(p, line)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)


# --- 内部辅助 ---

def _add_formatted_text(paragraph, text: str) -> None:
    """解析并添加含粗体/斜体的文本到段落"""
    # 正则匹配 **粗体** 和 *斜体*
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")
    last_end = 0
    for match in pattern.finditer(text):
        # 前面的普通文本
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        # 粗体或斜体
        if match.group(2):  # **粗体**
            paragraph.add_run(match.group(2)).bold = True
        elif match.group(3):  # *斜体*
            paragraph.add_run(match.group(3)).italic = True
        last_end = match.end()

    # 剩余普通文本
    if last_end < len(text):
        paragraph.add_run(text[last_end:])
