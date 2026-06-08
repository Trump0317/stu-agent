#!/usr/bin/env python3
"""Markdown → Word 转换

用法:
    python scripts/md_to_docx.py <input.md> <output.docx>

标题/段落/列表 → python-docx 新建文件。
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def convert(md_path: Path, output_path: Path):
    """转换 Markdown 为 .docx"""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题
        m = re.match(r"^(#{1,3})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        # 无序列表
        m = re.match(r"^[-*+]\s+(.+)", line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Bullet")
            i += 1
            continue

        # 有序列表
        m = re.match(r"^\d+\.\s+(.+)", line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            i += 1
            continue

        # 空行跳过
        if line.strip() == "":
            i += 1
            continue

        # 普通段落
        doc.add_paragraph(line.strip())
        i += 1

    doc.save(str(output_path))
    print(f"转换完成 → {output_path}")


def main():
    if len(sys.argv) < 3:
        print("用法: python md_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)

    md_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not md_path.exists():
        print(f"错误: 文件不存在: {md_path}", file=sys.stderr)
        sys.exit(1)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    convert(md_path, output_path)


if __name__ == "__main__":
    main()
