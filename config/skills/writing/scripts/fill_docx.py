#!/usr/bin/env python3
""".docx 模板填充工具

用法:
    python scripts/fill_docx.py <template.docx> <content.md> <output.docx>

按位置匹配：第 N 个 Heading ← 第 N 个 `##` 章节。
保留原模板样式、图片、表格、页眉页脚。
"""

import re
import sys
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


def split_md_sections(md_path: Path) -> list[str]:
    """按 ## 分裂 Markdown 为章节内容列表"""
    text = md_path.read_text(encoding="utf-8")
    # 按 ## 标题分割，跳过第一个（标题前的内容）
    parts = re.split(r"\n(?=## )", text)
    sections = []
    for part in parts:
        part = part.strip()
        if part and part.startswith("##"):
            sections.append(part)
    return sections


def insert_text_after(para, text: str, doc: Document):
    """在段落之后插入新段落，内容为 text，样式继承 Normal"""
    # 复制当前段落结构作为模板
    new_p_elem = deepcopy(para._element)
    # 清空原有 runs
    for r in new_p_elem.findall(qn("w:r")):
        new_p_elem.remove(r)
    # 创建新 run 写入文字
    r_elem = etree.SubElement(new_p_elem, qn("w:r"))
    t_elem = etree.SubElement(r_elem, qn("w:t"))
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    # 修改样式为 Normal（继承模板正文字体）
    pPr = new_p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(new_p_elem, qn("w:pPr"))
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "Normal")
    # 插入
    para._element.addnext(new_p_elem)


def fill_docx(template_path: Path, md_path: Path, output_path: Path):
    """填充 .docx 模板"""
    sections = split_md_sections(md_path)

    doc = Document(str(template_path))

    # 收集所有 Heading 段落
    heading_paras = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    if len(heading_paras) != len(sections):
        print(
            f"错误: 章节数不匹配 — 模板有 {len(heading_paras)} 个 Heading，"
            f"内容有 {len(sections)} 个 ## 章节",
            file=sys.stderr,
        )
        sys.exit(1)

    # 按位置匹配填充
    for heading_para, section_md in zip(heading_paras, sections):
        # 去掉 ## 标题行，取正文内容
        body_lines = section_md.splitlines()
        content = "\n".join(line for line in body_lines if not line.startswith("##")).strip()
        if content:
            for line in content.split("\n"):
                insert_text_after(heading_para, line, doc)
                # 更新 heading_para 指向刚插入的段落（后续插入紧跟）
                # 找到最后一个段落
                last = doc.paragraphs[-1]
                heading_para = last

    doc.save(str(output_path))
    print(f"填充完成 → {output_path}")


def main():
    if len(sys.argv) < 4:
        print("用法: python fill_docx.py <template.docx> <content.md> <output.docx>", file=sys.stderr)
        sys.exit(1)

    template_path = Path(sys.argv[1])
    md_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])

    if not template_path.exists():
        print(f"错误: 模板不存在: {template_path}", file=sys.stderr)
        sys.exit(1)
    if not md_path.exists():
        print(f"错误: 内容不存在: {md_path}", file=sys.stderr)
        sys.exit(1)

    fill_docx(template_path, md_path, output_path)


if __name__ == "__main__":
    main()
