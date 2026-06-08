"""Markdown → Word 转换单元测试

审核人: Trump
审核日期: 2026-06-08
审核状态: [已通过]
"""

import pytest
from docx import Document
from src.converter.markdown_to_word import convert_markdown_to_word


class TestConvertMarkdownToWord:
    def test_basic_conversion(self, tmp_path):
        """基本 Markdown → Word 转换，含 title"""
        md = "# 标题\n\n这是段落内容。"
        output = tmp_path / "output.docx"
        path = convert_markdown_to_word(md, str(output), title="测试文档")
        assert path == str(output)
        assert output.exists()
        assert output.stat().st_size > 0

        doc = Document(output)
        texts = [p.text for p in doc.paragraphs]
        assert "标题" in texts
        assert "这是段落内容。" in texts
        # title 作为文档第一行
        assert doc.paragraphs[0].text == "测试文档"

    def test_heading_levels(self, tmp_path):
        """标题层级正确（h1/h2/h3）"""
        md = "# 一级\n## 二级\n### 三级\n普通段落"
        output = tmp_path / "output.docx"
        convert_markdown_to_word(md, str(output))

        doc = Document(output)
        headings = [(p.style.name, p.text) for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert ("Heading 1", "一级") in headings
        assert ("Heading 2", "二级") in headings
        assert ("Heading 3", "三级") in headings

    def test_bullet_list(self, tmp_path):
        """无序列表 → Word bullet list"""
        md = "- 第一项\n- 第二项\n- 第三项"
        output = tmp_path / "output.docx"
        convert_markdown_to_word(md, str(output))

        doc = Document(output)
        bullet_texts = [p.text for p in doc.paragraphs if "List Bullet" in p.style.name]
        assert len(bullet_texts) == 3
        assert "第一项" in bullet_texts

    def test_bold_and_italic(self, tmp_path):
        """粗体/斜体保留"""
        md = "这是**粗体**和*斜体*文本"
        output = tmp_path / "output.docx"
        convert_markdown_to_word(md, str(output))

        doc = Document(output)
        para = doc.paragraphs[0]
        assert any(r.bold for r in para.runs)
        assert any(r.italic for r in para.runs)

    def test_mixed_content(self, tmp_path):
        """混合：标题 + 段落 + 列表 + 粗体"""
        md = "# 一级标题\n\n这是**粗体**段落。\n\n- 列表项一\n- 列表项二"
        output = tmp_path / "output.docx"
        convert_markdown_to_word(md, str(output))

        doc = Document(output)
        texts = [p.text for p in doc.paragraphs]
        assert "一级标题" in texts
        assert "列表项一" in texts
        assert any(r.bold for r in doc.paragraphs[1].runs)

    def test_chinese_content(self, tmp_path):
        """中文长文本（含标点）"""
        md = "# 课程论文\n\n本研究旨在探讨人工智能对高等教育的影响。\n\n## 引言\n\n近年来，AI技术快速发展，「深度学习」与《大语言模型》成为热点。"
        output = tmp_path / "output.docx"
        convert_markdown_to_word(md, str(output))

        doc = Document(output)
        full_text = " ".join([p.text for p in doc.paragraphs])
        assert "课程论文" in full_text
        assert "深度学习" in full_text

    def test_empty_content_raises(self, tmp_path):
        """空内容抛 ValueError"""
        with pytest.raises(ValueError, match="content 不能为空"):
            convert_markdown_to_word("", str(tmp_path / "empty.docx"))

    def test_output_overwrites_existing(self, tmp_path):
        """已存在文件被覆盖"""
        output = tmp_path / "output.docx"
        output.write_text("old")
        convert_markdown_to_word("# 新内容", str(output))
        doc = Document(output)
        assert doc.paragraphs[0].text == "新内容"
        assert "old" not in [p.text for p in doc.paragraphs]
