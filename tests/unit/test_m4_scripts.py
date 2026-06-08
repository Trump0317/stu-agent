"""M4-3 Skill 脚本单元测试

测试 writing skill 的 6 个脚本：read_template, fill_docx, md_to_docx, md_to_pdf, docx_to_pdf, format

审核人: Trump
审核日期: 2026-06-08
审核状态: [已通过]
"""

import json
import os
import subprocess
import sys
from pathlib import Path

import pytest

SKILL_SCRIPTS = Path(__file__).parent.parent.parent / "config" / "skills" / "writing" / "scripts"


def _run(script: str, *args: str) -> subprocess.CompletedProcess:
    """运行 Skill 脚本"""
    script_path = SKILL_SCRIPTS / script
    if not script_path.exists():
        pytest.fail(f"脚本不存在: {script_path}")
    cmd = [sys.executable, str(script_path), *args]
    return subprocess.run(cmd, capture_output=True, text=True, timeout=30)


def _make_md(tmp_path, name="test.md", content="") -> Path:
    p = tmp_path / name
    p.write_text(content, encoding="utf-8")
    return p


def _make_yaml(tmp_path, name="test.yaml", sections=None) -> Path:
    if sections is None:
        sections = [{"title": "引言"}, {"title": "正文"}, {"title": "结论"}]
    yaml_content = f"name: 测试模板\ntype: test\nsections:\n"
    for s in sections:
        yaml_content += f"  - title: {s['title']}\n"
    p = tmp_path / name
    p.write_text(yaml_content, encoding="utf-8")
    return p


def _make_docx(tmp_path, name="test.docx", headings=None) -> Path:
    """创建含 Heading 的测试 .docx"""
    from docx import Document
    if headings is None:
        headings = ["第一章 引言", "第二章 正文", "第三章 结论"]
    doc = Document()
    for h in headings:
        doc.add_heading(h, level=1)
        doc.add_paragraph(f"{h}的占位内容。")
    p = tmp_path / name
    doc.save(str(p))
    return p


# ============================================================
# read_template.py
# ============================================================

class TestReadTemplate:
    """M4-1: read_template.py

    [~]
    """

    def test_read_yaml(self, tmp_path):
        """读取 .yaml 模板，返回章节结构"""
        yaml_path = _make_yaml(tmp_path)
        r = _run("read_template.py", str(yaml_path))
        assert r.returncode == 0, r.stderr
        data = json.loads(r.stdout)
        assert data["name"] == "测试模板"
        assert [s["title"] for s in data["sections"]] == ["引言", "正文", "结论"]

    def test_read_md(self, tmp_path):
        """读取 .md 模板，提取 #/## 为章节"""
        md_path = _make_md(tmp_path, content=(
            "# 课程论文\n\n"
            "## 摘要\n摘要内容\n\n"
            "## 引言\n引言内容\n\n"
            "## 结论\n结论内容\n"
        ))
        r = _run("read_template.py", str(md_path))
        assert r.returncode == 0, r.stderr
        data = json.loads(r.stdout)
        assert "摘要" in [s["title"] for s in data["sections"]]
        assert "引言" in [s["title"] for s in data["sections"]]
        assert "结论" in [s["title"] for s in data["sections"]]
        # # 一级标题作为文档名，不应出现在 sections 中
        assert data.get("name") == "课程论文" or "课程论文" not in [s["title"] for s in data["sections"]]

    def test_read_docx(self, tmp_path):
        """读取 .docx 模板，提取 Heading 为章节"""
        docx_path = _make_docx(tmp_path)
        r = _run("read_template.py", str(docx_path))
        assert r.returncode == 0, r.stderr
        data = json.loads(r.stdout)
        titles = [s["title"] for s in data["sections"]]
        assert "第一章 引言" in titles
        assert "第二章 正文" in titles
        assert "第三章 结论" in titles

    def test_file_not_found(self, tmp_path):
        """文件不存在时返回非零退出码"""
        r = _run("read_template.py", str(tmp_path / "nonexistent.yaml"))
        assert r.returncode != 0

    def test_unsupported_format(self, tmp_path):
        """不支持的格式返回错误"""
        bad = tmp_path / "test.txt"
        bad.write_text("hello")
        r = _run("read_template.py", str(bad))
        assert r.returncode != 0


# ============================================================
# fill_docx.py
# ============================================================

class TestFillDocx:
    """M4-1: fill_docx.py — 填充 .docx，保留样式

    [~]
    """

    def test_fill_by_position(self, tmp_path):
        """按位置匹配：第 N 个 Heading ← 第 N 个 ## 章节（含多行内容）"""
        docx_path = _make_docx(tmp_path, headings=["第一章 引言", "第二章 正文"])
        md_path = _make_md(tmp_path, "content.md", (
            "## 引言\n\n这是引言第一行。\n这是引言第二行。\n\n"
            "## 正文\n\n这是正文内容。\n"
        ))
        out = tmp_path / "output.docx"

        r = _run("fill_docx.py", str(docx_path), str(md_path), str(out))
        assert r.returncode == 0, r.stderr
        assert out.exists()

        from docx import Document
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        # 多行内容都在第一章 heading 之后
        ch1_idx = texts.index("第一章 引言")
        ch2_idx = texts.index("第二章 正文")
        assert "这是引言第一行。" in texts[ch1_idx + 1 : ch2_idx]
        assert "这是引言第二行。" in texts[ch1_idx + 1 : ch2_idx]
        assert "这是正文内容。" in texts[ch2_idx + 1 :]

    def test_tables_preserved(self, tmp_path):
        """填充后表格保留"""
        from docx import Document
        docx_path = _make_docx(tmp_path, headings=["第一章"])
        # 在模板中插入表格
        doc = Document(str(docx_path))
        doc.add_table(rows=2, cols=2)
        doc.save(str(docx_path))

        md_path = _make_md(tmp_path, "content.md", "## 第一章\n\n内容。\n")
        out = tmp_path / "output.docx"

        r = _run("fill_docx.py", str(docx_path), str(md_path), str(out))
        assert r.returncode == 0

        doc2 = Document(str(out))
        assert len(doc2.tables) == 1, "表格应保留"

    def test_section_count_mismatch(self, tmp_path):
        """章节数量不匹配时返回错误"""
        docx_path = _make_docx(tmp_path, headings=["第一章", "第二章", "第三章"])
        md_path = _make_md(tmp_path, "content.md", "## 第一章\n\n内容。\n")
        out = tmp_path / "output.docx"

        r = _run("fill_docx.py", str(docx_path), str(md_path), str(out))
        assert r.returncode != 0, "3 个 Heading vs 1 个 ## 章节应报错"
        assert "不匹配" in (r.stdout + r.stderr) or "mismatch" in (r.stdout + r.stderr).lower()

    def test_file_not_found(self, tmp_path):
        """模板不存在时报错"""
        out = tmp_path / "out.docx"
        r = _run("fill_docx.py", str(tmp_path / "no.docx"), str(tmp_path / "no.md"), str(out))
        assert r.returncode != 0


# ============================================================
# md_to_docx.py
# ============================================================

class TestMdToDocx:
    """M4-1: md_to_docx.py — Markdown → .docx 新建

    [~]
    """

    def test_basic_conversion(self, tmp_path):
        """标题、段落、列表正确转换"""
        md_path = _make_md(tmp_path, content=(
            "# 文档标题\n\n"
            "## 第一章\n\n"
            "这是第一段。\n\n"
            "这是第二段。\n\n"
            "- 列表项 1\n"
            "- 列表项 2\n\n"
            "## 第二章\n\n"
            "第二章内容。\n"
        ))
        out = tmp_path / "output.docx"

        r = _run("md_to_docx.py", str(md_path), str(out))
        assert r.returncode == 0, r.stderr
        assert out.exists()

        from docx import Document
        doc = Document(str(out))
        styles = [(p.style.name, p.text[:30]) for p in doc.paragraphs]
        # 验证标题层级
        heading_texts = [t for s, t in styles if "Heading" in s]
        assert "第一章" in heading_texts
        assert "第二章" in heading_texts

    def test_output_file_created(self, tmp_path):
        """输出文件存在且非空"""
        md_path = _make_md(tmp_path, content="# Hello\n\nWorld.\n")
        out = tmp_path / "out.docx"

        r = _run("md_to_docx.py", str(md_path), str(out))
        assert r.returncode == 0
        assert out.exists()
        assert out.stat().st_size > 0

    def test_file_not_found(self, tmp_path):
        """输入文件不存在时报错"""
        r = _run("md_to_docx.py", str(tmp_path / "no.md"), str(tmp_path / "out.docx"))
        assert r.returncode != 0


# ============================================================
# md_to_pdf.py
# ============================================================

class TestMdToPdf:
    """M4-1: md_to_pdf.py — Markdown → PDF

    [~]
    """

    def test_output_pdf_created(self, tmp_path):
        """Markdown 可转为 PDF"""
        md_path = _make_md(tmp_path, content="# 测试\n\n这是内容。\n")
        out = tmp_path / "output.pdf"

        r = _run("md_to_pdf.py", str(md_path), str(out))
        # PDF 生成可能依赖 weasyprint 或 LibreOffice，未安装时跳过
        if r.returncode != 0:
            if "No module named" in r.stderr or "libreoffice" in r.stderr.lower():
                pytest.skip("PDF 依赖未安装")
        assert r.returncode == 0, r.stderr
        assert out.exists()
        assert out.stat().st_size > 0


# ============================================================
# docx_to_pdf.py
# ============================================================

class TestDocxToPdf:
    """M4-1: docx_to_pdf.py — .docx → PDF

    [~]
    """

    def test_docx_to_pdf(self, tmp_path):
        """.docx 转为 PDF"""
        docx_path = _make_docx(tmp_path, headings=["测试文档"])
        out = tmp_path / "output.pdf"

        r = _run("docx_to_pdf.py", str(docx_path), str(out))
        if r.returncode != 0:
            if "libreoffice" in r.stderr.lower():
                pytest.skip("LibreOffice 未安装")
        assert r.returncode == 0, r.stderr
        assert out.exists()
        assert out.stat().st_size > 0


# ============================================================
# format.py
# ============================================================

class TestFormat:
    """M4-1: format.py — Markdown 格式化

    [~]
    """

    def test_clean_trailing_spaces(self, tmp_path):
        """清理行尾空格"""
        md = _make_md(tmp_path, content="hello   \nworld\n")
        out = tmp_path / "out.md"
        r = _run("format.py", str(md), str(out))
        assert r.returncode == 0
        content = out.read_text()
        assert "   " not in content

    def test_merge_blank_lines(self, tmp_path):
        """合并多余空行"""
        md = _make_md(tmp_path, content="line1\n\n\n\nline2\n")
        out = tmp_path / "out.md"
        r = _run("format.py", str(md), str(out))
        assert r.returncode == 0
        content = out.read_text()
        assert "\n\n\n\n" not in content

    def test_inplace_edit(self, tmp_path):
        """不指定输出时原地修改"""
        md = _make_md(tmp_path, content="hello   \n\n\n\nworld\n")
        r = _run("format.py", str(md))
        assert r.returncode == 0
        content = md.read_text()
        assert "   " not in content
        assert "\n\n\n\n" not in content

    def test_heading_blank_line(self, tmp_path):
        """标题前自动补空行"""
        md = _make_md(tmp_path, content="正文内容\n# 标题\n正文\n")
        out = tmp_path / "out.md"
        r = _run("format.py", str(md), str(out))
        assert r.returncode == 0
        content = out.read_text()
        # 标题前应有空行
        assert "\n\n# 标题" in content
